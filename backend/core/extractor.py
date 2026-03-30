"""
Phase 3: Text extraction pipeline.
Handles PDF, DOCX, TXT, and MD files with robust error handling.
Returns extracted text and page boundary metadata.
"""

import io
import logging
from pathlib import Path
from dataclasses import dataclass
from typing import Optional

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
MAX_TEXT_LENGTH = 500_000  # ~125K tokens max

# Magic bytes for binary formats
_MAGIC = {
    ".pdf":  b"%PDF",
    ".docx": b"PK\x03\x04",  # ZIP container
}


@dataclass
class ExtractionResult:
    text: str
    page_count: int = 1
    page_boundaries: Optional[list[int]] = None
    file_type: str = ""
    warnings: list[str] = None

    def __post_init__(self):
        self.warnings = self.warnings or []


class TextExtractor:
    """Extracts text content from various document formats."""

    def extract(self, content_bytes: bytes, filename: str) -> ExtractionResult:
        ext = Path(filename).suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}")

        extractors = {
            ".txt": self._extract_txt,
            ".md": self._extract_txt,
            ".pdf": self._extract_pdf,
            ".docx": self._extract_docx,
        }

        # Verify magic bytes for binary formats
        expected_magic = _MAGIC.get(ext)
        if expected_magic and not content_bytes.startswith(expected_magic):
            raise ValueError(
                f"File content does not match the declared type '{ext}'. "
                "Please upload a valid file."
            )

        result = extractors[ext](content_bytes)
        result.file_type = ext.lstrip(".")

        if not result.text.strip():
            raise ValueError(
                f"No text could be extracted from '{filename}'. The file may be empty or image-based.")

        if len(result.text) > MAX_TEXT_LENGTH:
            result.text = result.text[:MAX_TEXT_LENGTH]
            result.warnings.append(
                f"Text truncated to {MAX_TEXT_LENGTH} characters")

        return result

    def _extract_txt(self, content_bytes: bytes) -> ExtractionResult:
        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                text = content_bytes.decode(encoding)
                return ExtractionResult(text=text, file_type="txt")
            except UnicodeDecodeError:
                continue
        raise ValueError(
            "Could not decode text file with any supported encoding")

    def _extract_pdf(self, content_bytes: bytes) -> ExtractionResult:
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError(
                "PyPDF2 is required for PDF processing. Install with: pip install PyPDF2")

        reader = PdfReader(io.BytesIO(content_bytes))
        texts = []
        page_boundaries = []
        char_pos = 0
        warnings = []

        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
                texts.append(text)
                char_pos += len(text) + 2
                page_boundaries.append(char_pos)
            except Exception as e:
                warnings.append(f"Page {i+1}: extraction failed ({e})")
                page_boundaries.append(char_pos)

        full_text = "\n\n".join(texts)

        return ExtractionResult(
            text=full_text,
            page_count=len(reader.pages),
            page_boundaries=page_boundaries,
            file_type="pdf",
            warnings=warnings,
        )

    def _extract_docx(self, content_bytes: bytes) -> ExtractionResult:
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX processing. Install with: pip install python-docx")

        doc = Document(io.BytesIO(content_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip()
                                      for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        return ExtractionResult(
            text="\n\n".join(paragraphs),
            page_count=1,
            file_type="docx",
        )
